"""
笔记附件正文抽取（上传与播客加载共用）。

- PDF：优先 PyMuPDF（fitz），失败或几乎无字时回退 PyPDF2。
- DOCX：优先 python-docx，失败回退 word/document.xml 正则。
- 纯文本：charset-normalizer 探测编码，失败再用 utf-8 ignore。
- HTML/HTM/XHTML：按块级 DOM 分段（段间双换行），与 EPUB 同源抽取逻辑。
- EPUB：临时文件 + content_parser.parse_epub（避免重复实现 spine 逻辑）。
- DOC：临时文件 + antiword / catdoc / soffice（与旧逻辑一致）。
- CSV / XLSX / XLS：表格线性化为 Markdown（含表头），大表按行列上限截断并注明；.xls 使用 xlrd。
- 图片（png/jpg/jpeg/webp/gif/avif）：笔记上传入口已不再接收；若库内仍有历史附件，可走可配置视觉模型 OCR（Qwen VL），未配置时仅存档。
"""
from __future__ import annotations

import base64
import csv
import io
import logging
import os
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from typing import Any

from .providers.openai_compat_text import chat_completion_openai_compatible

logger = logging.getLogger(__name__)

# 表格导入上限（防极端宽表拖垮内存与向量）
_MAX_CSV_ROWS = 50_000
_MAX_CSV_FIELD = 12_000
_MAX_XLSX_SHEETS = 15
_MAX_XLSX_ROWS_PER_SHEET = 800
_MAX_XLSX_COLS = 48


def _md_cell(s: str) -> str:
    t = (s or "").replace("\r\n", "\n").replace("\r", "\n").replace("|", "\\|").strip()
    t = re.sub(r"\s+", " ", t)
    return t[:800]


def _csv_bytes_to_text(data: bytes) -> NoteParseResult:
    if not data:
        return NoteParseResult(text="", status="empty", engine="csv", detail="空文件")
    text, enc = _decode_plain_bytes(data)
    if not text.strip():
        return NoteParseResult(text="", status="empty", engine="csv", detail="CSV 解码后无内容")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except Exception:
        dialect = csv.excel
    lines_out: list[str] = []
    try:
        reader = csv.reader(io.StringIO(text), dialect=dialect)
        for i, row in enumerate(reader):
            if i >= _MAX_CSV_ROWS:
                lines_out.append(f"\n（CSV 已截断，仅保留前 {_MAX_CSV_ROWS} 行）")
                break
            cells = [_md_cell(c)[:_MAX_CSV_FIELD] for c in row]
            if any(cells):
                lines_out.append("| " + " | ".join(cells) + " |")
    except Exception as exc:
        return NoteParseResult(text="", status="error", engine="csv", detail=f"CSV 解析失败：{exc}"[:400])
    if not lines_out:
        return NoteParseResult(text="", status="empty", engine="csv", detail="CSV 无有效行")
    header = lines_out[0]
    sep_cells = ["---"] * max(1, header.count("|") - 1)
    sep = "| " + " | ".join(sep_cells) + " |"
    md = "\n".join([header, sep, *lines_out[1:]])
    md = md.strip()
    seg = [{"text": md, "meta": {"block_type": "csv", "heading_path": []}}]
    return NoteParseResult(
        text=md, status="ok", engine=f"csv:{enc or 'utf-8'}", detail=None, encoding=enc, rag_segments=seg
    )


def _xlsx_bytes_to_text(data: bytes) -> NoteParseResult:
    if not data:
        return NoteParseResult(text="", status="empty", engine="xlsx", detail="空文件")
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        return NoteParseResult(
            text="",
            status="error",
            engine="xlsx",
            detail=f"未安装 openpyxl：{exc}"[:200],
        )
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:
        return NoteParseResult(
            text="",
            status="error",
            engine="xlsx",
            detail=f"无法打开 XLSX（可能已加密或损坏）：{exc}"[:400],
        )
    parts: list[str] = []
    try:
        sheet_names = wb.sheetnames[:_MAX_XLSX_SHEETS]
        for sname in sheet_names:
            ws = wb[sname]
            parts.append(f"## 工作表: {_md_cell(sname)}")
            rows_iter = ws.iter_rows(max_row=_MAX_XLSX_ROWS_PER_SHEET, max_col=_MAX_XLSX_COLS, values_only=True)
            grid: list[list[str]] = []
            for row in rows_iter:
                cells = [_md_cell(str(c) if c is not None else "") for c in row]
                if any(x.strip() for x in cells):
                    grid.append(cells)
            if not grid:
                parts.append("（本工作表无有效单元格）")
                continue
            header = grid[0]
            parts.append("| " + " | ".join(header) + " |")
            parts.append("| " + " | ".join(["---"] * len(header)) + " |")
            for r in grid[1:]:
                if len(r) < len(header):
                    r = r + [""] * (len(header) - len(r))
                parts.append("| " + " | ".join(r[: len(header)]) + " |")
            if getattr(ws, "max_row", 0) > _MAX_XLSX_ROWS_PER_SHEET:
                parts.append(f"（本表已截断，仅保留前 {_MAX_XLSX_ROWS_PER_SHEET} 行）")
            parts.append("")
    finally:
        try:
            wb.close()
        except Exception:
            pass
    out = "\n".join(parts).strip()
    if not out:
        return NoteParseResult(text="", status="empty", engine="xlsx", detail="XLSX 未解析出单元格文本")
    xlsx_segments: list[dict[str, Any]] = []
    for part in re.split(r"(?m)^(?=## 工作表:\s)", out):
        part = (part or "").strip()
        if not part:
            continue
        m_sheet = re.match(r"## 工作表:\s*(.+?)(?:\n|$)", part)
        sheet = (m_sheet.group(1).strip() if m_sheet else "") or ""
        body = part[m_sheet.end() :].strip() if m_sheet else part
        if body:
            xlsx_segments.append(
                {
                    "text": body,
                    "meta": {
                        "block_type": "sheet",
                        "sheet": sheet,
                        "heading_path": [sheet] if sheet else [],
                    },
                }
            )
    if not xlsx_segments:
        xlsx_segments = [{"text": out, "meta": {"block_type": "xlsx", "heading_path": []}}]
    return NoteParseResult(text=out, status="ok", engine="xlsx-openpyxl", detail=None, rag_segments=xlsx_segments)


def _xls_bytes_to_text(data: bytes) -> NoteParseResult:
    """经典 Excel 97-2003（.xls），使用 xlrd 读表。"""
    if not data:
        return NoteParseResult(text="", status="empty", engine="xls", detail="空文件")
    try:
        import xlrd  # type: ignore
    except Exception as exc:
        return NoteParseResult(
            text="",
            status="error",
            engine="xls",
            detail=f"未安装 xlrd：{exc}"[:200],
        )
    try:
        book = xlrd.open_workbook(file_contents=data, formatting_info=False)
    except Exception as exc:
        return NoteParseResult(
            text="",
            status="error",
            engine="xls",
            detail=f"无法打开 XLS（可能已损坏或非 Excel 表格）：{exc}"[:400],
        )
    parts: list[str] = []
    try:
        n_sheets = min(book.nsheets, _MAX_XLSX_SHEETS)
        for si in range(n_sheets):
            sh = book.sheet_by_index(si)
            sname = str(sh.name or f"Sheet{si + 1}")
            parts.append(f"## 工作表: {_md_cell(sname)}")
            nrows = min(sh.nrows, _MAX_XLSX_ROWS_PER_SHEET)
            ncols = min(sh.ncols, _MAX_XLSX_COLS) if sh.ncols else 0
            grid: list[list[str]] = []
            for ri in range(nrows):
                row_cells: list[str] = []
                for ci in range(ncols):
                    try:
                        val = sh.cell_value(ri, ci)
                        row_cells.append(_md_cell(str(val) if val is not None else ""))
                    except Exception:
                        row_cells.append("")
                if any(x.strip() for x in row_cells):
                    grid.append(row_cells)
            if not grid:
                parts.append("（本工作表无有效单元格）")
                continue
            header = grid[0]
            parts.append("| " + " | ".join(header) + " |")
            parts.append("| " + " | ".join(["---"] * len(header)) + " |")
            for r in grid[1:]:
                if len(r) < len(header):
                    r = r + [""] * (len(header) - len(r))
                parts.append("| " + " | ".join(r[: len(header)]) + " |")
            if sh.nrows > _MAX_XLSX_ROWS_PER_SHEET:
                parts.append(f"（本表已截断，仅保留前 {_MAX_XLSX_ROWS_PER_SHEET} 行）")
            parts.append("")
    finally:
        try:
            book.release_resources()
        except Exception:
            pass
    out = "\n".join(parts).strip()
    if not out:
        return NoteParseResult(text="", status="empty", engine="xls", detail="XLS 未解析出单元格文本")
    xls_segments: list[dict[str, Any]] = []
    for part in re.split(r"(?m)^(?=## 工作表:\s)", out):
        part = (part or "").strip()
        if not part:
            continue
        m_sheet = re.match(r"## 工作表:\s*(.+?)(?:\n|$)", part)
        sheet = (m_sheet.group(1).strip() if m_sheet else "") or ""
        body = part[m_sheet.end() :].strip() if m_sheet else part
        if body:
            xls_segments.append(
                {
                    "text": body,
                    "meta": {
                        "block_type": "sheet",
                        "sheet": sheet,
                        "heading_path": [sheet] if sheet else [],
                    },
                }
            )
    if not xls_segments:
        xls_segments = [{"text": out, "meta": {"block_type": "xls", "heading_path": []}}]
    return NoteParseResult(text=out, status="ok", engine="xls-xlrd", detail=None, rag_segments=xls_segments)


@dataclass
class NoteParseResult:
    """单次解析结果，供写入 metadata 与 API 返回。"""

    text: str
    status: str  # ok | empty | error
    engine: str
    detail: str | None = None
    encoding: str | None = None
    # 结构化分段（可选）：用于 RAG chunk_meta（heading_path / 表格 / PDF 页码等）
    rag_segments: list[dict[str, Any]] | None = None

    @property
    def ok(self) -> bool:
        return bool((self.text or "").strip())


def _docx_xml_fallback(data: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            xml_data = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"</w:p>", "\n", xml_data)
        text = re.sub(r"<[^>]+>", "", text)
        return re.sub(r"\n{2,}", "\n", text).strip()
    except Exception:
        return ""


def _docx_python_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue
            st_name = ""
            try:
                st_name = str((p.style.name if p.style else "") or "")
            except Exception:
                st_name = ""
            if st_name.startswith("Heading"):
                try:
                    lev = int(st_name.replace("Heading", "").strip() or "2")
                except Exception:
                    lev = 2
                lev = max(1, min(6, lev))
                lines.append("#" * lev + " " + t)
            elif "标题" in st_name and not st_name.startswith("Heading"):
                lines.append("## " + t)
            else:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
                if cells:
                    lines.append("\t".join(cells))
        return "\n\n".join(lines).strip()
    except Exception as exc:
        logger.debug("python-docx parse failed: %s", exc)
        return ""


def _rows_to_markdown_table(rows_txt: list[str]) -> str:
    """将 TSV/制表符行转为 Markdown 表，便于 rag_core 按行切块。"""
    grid: list[list[str]] = []
    for line in rows_txt:
        cells = [c.strip() for c in line.split("\t") if c is not None]
        if not cells:
            cells = [line.strip()]
        if any(cells):
            grid.append([c.replace("|", "\\|") for c in cells])
    if not grid:
        return ""
    ncol = max(len(r) for r in grid)
    norm = [r + [""] * (ncol - len(r)) for r in grid]
    header = norm[0]
    sep = ["---"] * ncol
    body = norm[1:] if len(norm) > 1 else []
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _docx_rag_segments(data: bytes) -> list[dict[str, Any]]:
    """按文档流顺序输出段落/表格块，带 heading_path。"""
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except Exception:
        return []

    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return []

    heading_path: list[str] = []
    segments: list[dict[str, Any]] = []

    try:
        for el in doc.element.body:
            if el.tag == qn("w:p"):
                p = Paragraph(el, doc)
                t = (p.text or "").strip()
                if not t:
                    continue
                st_name = ""
                try:
                    st_name = str((p.style.name if p.style else "") or "")
                except Exception:
                    pass
                if st_name.startswith("Heading"):
                    try:
                        lev = int(st_name.replace("Heading", "").strip() or "2")
                    except Exception:
                        lev = 2
                    lev = max(1, min(6, lev))
                    heading_path = heading_path[: lev - 1] + [t]
                    segments.append(
                        {
                            "text": t,
                            "meta": {
                                "block_type": "heading",
                                "level": lev,
                                "heading_path": list(heading_path),
                            },
                        }
                    )
                elif "标题" in st_name and not st_name.startswith("Heading"):
                    heading_path = heading_path[:1] + [t] if heading_path else [t]
                    segments.append(
                        {
                            "text": t,
                            "meta": {
                                "block_type": "heading",
                                "level": 2,
                                "heading_path": list(heading_path),
                            },
                        }
                    )
                else:
                    segments.append(
                        {"text": t, "meta": {"block_type": "paragraph", "heading_path": list(heading_path)}}
                    )
            elif el.tag == qn("w:tbl"):
                table = Table(el, doc)
                rows_txt: list[str] = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
                    if cells:
                        rows_txt.append("\t".join(cells))
                if rows_txt:
                    tbl = _rows_to_markdown_table(rows_txt)
                    if not tbl:
                        tbl = "\n".join(rows_txt)
                    segments.append(
                        {
                            "text": tbl,
                            "meta": {
                                "block_type": "table",
                                "heading_path": list(heading_path),
                                "table_rows": len(rows_txt),
                            },
                        }
                    )
    except Exception as exc:
        logger.debug("docx rag segments failed: %s", exc)
        return []
    return segments


def _normalize_prose_plaintext(text: str) -> str:
    """将 antiword 等输出的纯文本规范为段间双换行，便于阅读器分块。"""
    raw = str(text or "").replace("\r\n", "\n").strip()
    if not raw:
        return ""
    if re.search(r"\n\s*\n", raw):
        parts = [p.strip() for p in re.split(r"\n\s*\n", raw) if p.strip()]
    else:
        parts = [p.strip() for p in raw.split("\n") if p.strip()]
    return "\n\n".join(parts)


def _plain_text_rag_segments(text: str) -> list[dict[str, Any]]:
    """按段/标题切分纯文本（DOC 等无 DOM 来源）。"""
    normalized = _normalize_prose_plaintext(text)
    if not normalized:
        return []
    heading_path: list[str] = []
    segments: list[dict[str, Any]] = []
    for para in normalized.split("\n\n"):
        p = para.strip()
        if not p:
            continue
        hm = re.match(r"^(#{1,3})\s+(.+)$", p)
        if hm:
            lev = len(hm.group(1))
            title = hm.group(2).strip()
            heading_path = heading_path[: lev - 1] + [title]
            segments.append(
                {
                    "text": title,
                    "meta": {
                        "block_type": "heading",
                        "level": lev,
                        "heading_path": list(heading_path),
                    },
                }
            )
            continue
        segments.append(
            {"text": p, "meta": {"block_type": "paragraph", "heading_path": list(heading_path)}}
        )
    return segments


def _pdf_rag_segments_by_page(data: bytes) -> list[dict[str, Any]]:
    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return []
    try:
        out: list[dict[str, Any]] = []
        for i in range(len(doc)):
            t = (doc[i].get_text() or "").strip()
            if t:
                out.append({"text": t, "meta": {"block_type": "pdf_page", "page": i + 1, "heading_path": []}})
        return out
    finally:
        try:
            doc.close()
        except Exception:
            pass


def _pdf_pymupdf_chars_per_page(data: bytes) -> tuple[int, int]:
    """(页数, 总字符数) — 用于扫描件启发式。"""
    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            n = len(doc)
            ch = 0
            for i in range(n):
                ch += len((doc[i].get_text() or ""))
            return n, ch
        finally:
            doc.close()
    except Exception:
        return 0, 0


def _decode_plain_bytes(data: bytes) -> tuple[str, str | None]:
    if not data:
        return "", None
    # 常见带 BOM 文本优先：避免 UTF-16/UTF-32 被误判后出现乱码或空文本。
    bom_codecs: list[tuple[bytes, str]] = [
        (b"\xef\xbb\xbf", "utf-8-sig"),
        (b"\xff\xfe\x00\x00", "utf-32-le"),
        (b"\x00\x00\xfe\xff", "utf-32-be"),
        (b"\xff\xfe", "utf-16-le"),
        (b"\xfe\xff", "utf-16-be"),
    ]
    for bom, enc in bom_codecs:
        if data.startswith(bom):
            try:
                txt = data.decode(enc)
                if txt.strip():
                    return txt, enc
            except Exception:
                break

    # 无 BOM 时按常见编码做一轮严格解码尝试（先 utf-8，再 utf-16，再 gb18030）。
    preferred_codecs = ("utf-8", "utf-16", "utf-16-le", "utf-16-be", "gb18030")
    for enc in preferred_codecs:
        try:
            txt = data.decode(enc)
            if txt.strip():
                return txt, enc
        except Exception:
            continue
    try:
        from charset_normalizer import from_bytes  # type: ignore

        best = from_bytes(data).best()
        if best is not None:
            s = str(best)
            enc = getattr(best, "encoding", None)
            if s.strip():
                return s, enc
    except Exception as exc:
        logger.debug("charset_normalizer failed: %s", exc)
    # 最后兜底：保证不抛错，并尽量保留可见字符。
    return data.decode("utf-8", errors="ignore"), "utf-8-ignore"


def _pdf_pymupdf(data: bytes) -> tuple[str, bool]:
    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            parts: list[str] = []
            for i in range(len(doc)):
                parts.append(doc[i].get_text() or "")
            text = "\n".join(parts).strip()
            return text, bool(text.strip())
        finally:
            doc.close()
    except Exception as exc:
        logger.debug("pymupdf failed: %s", exc)
        return "", False


def _pdf_pypdf2(data: bytes) -> tuple[str, bool]:
    try:
        from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            except Exception:
                continue
        text = "\n".join(parts).strip()
        return text, bool(text.strip())
    except Exception as exc:
        logger.debug("pypdf2 failed: %s", exc)
        return "", False


def _epub_via_content_parser(path: str) -> tuple[str, bool, list[dict[str, Any]] | None]:
    try:
        from app.fyv_shared.content_parser import content_parser

        r = content_parser.parse_epub(path)
        if r.get("success"):
            text = str(r.get("content") or "").strip()
            raw_segs = r.get("rag_segments")
            segs = raw_segs if isinstance(raw_segs, list) and raw_segs else None
            return text, True, segs
        return "", False, None
    except Exception as exc:
        logger.warning("epub parse: %s", exc)
        return "", False, None


def _parse_doc_path(path: str) -> tuple[str, bool]:
    candidates: list[tuple[list[str], str]] = [
        (["antiword", path], "antiword"),
        (["catdoc", path], "catdoc"),
    ]
    for cmd, _name in candidates:
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if proc.returncode == 0:
                txt = (proc.stdout or "").strip()
                if txt:
                    return txt, True
        except Exception:
            continue
    try:
        out_dir = os.path.dirname(path)
        proc = subprocess.run(
            ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", out_dir, path],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if proc.returncode == 0:
            txt_path = os.path.splitext(path)[0] + ".txt"
            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                    txt = f.read().strip()
                try:
                    os.remove(txt_path)
                except OSError:
                    pass
                if txt:
                    return txt, True
    except Exception:
        pass
    return "", False


def _extract_html_from_bytes(data: bytes) -> NoteParseResult:
    """从本地 HTML/XHTML 按块级 DOM 抽取正文（段间 \\n\\n），对齐 EPUB 阅读器分块。"""
    if not data:
        return NoteParseResult(text="", status="empty", engine="html", detail="空文件")
    try:
        from bs4 import BeautifulSoup

        from .fyv_shared.content_parser import (
            _extract_epub_blocks_from_soup,
            _html_blocks_to_markdown_and_segments,
        )
    except Exception as exc:  # pragma: no cover
        logger.warning("HTML parse deps unavailable: %s", exc)
        return NoteParseResult(
            text="",
            status="error",
            engine="html",
            detail="HTML 解析依赖不可用",
        )
    raw, _enc = _decode_plain_bytes(data)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style", "noscript", "template"]):
        tag.decompose()
    for tag in soup.find_all(["iframe", "object", "embed", "video", "audio", "svg", "canvas"]):
        tag.decompose()
    html_blocks = _extract_epub_blocks_from_soup(soup)
    if not html_blocks:
        text_out = soup.get_text(separator="\n", strip=True)
        lines = [ln.strip() for ln in text_out.split("\n") if ln.strip()]
        content = _normalize_prose_plaintext("\n".join(lines))
        segs = _plain_text_rag_segments(content) if content else []
        if content:
            return NoteParseResult(
                text=content,
                status="ok",
                engine="html-bs4-fallback",
                detail=None,
                rag_segments=segs if segs else None,
            )
        return NoteParseResult(
            text="",
            status="empty",
            engine="html-bs4",
            detail="HTML 中未解析出可见文本（可能为框架页或需登录内容）",
        )
    content, segs = _html_blocks_to_markdown_and_segments(html_blocks, chapter_mode=False)
    if content.strip():
        return NoteParseResult(
            text=content.strip(),
            status="ok",
            engine="html-bs4-blocks",
            detail=None,
            rag_segments=segs if segs else None,
        )
    return NoteParseResult(
        text="",
        status="empty",
        engine="html-bs4",
        detail="HTML 中未解析出可见文本（可能为框架页或需登录内容）",
    )


def _ocr_image_via_openai_compat(data: bytes, ext: str) -> NoteParseResult:
    """
    图片 OCR（可选能力）：
    - 需配置 QWEN_API_KEY + QWEN_BASE_URL
    - 模型默认 QWEN_VL_MODEL=qwen-vl-plus（可覆写）
    """
    key = str(os.getenv("QWEN_API_KEY") or "").strip()
    base = str(os.getenv("QWEN_BASE_URL") or "").strip()
    model = str(os.getenv("QWEN_VL_MODEL") or "qwen-vl-plus").strip()
    if not key or not base:
        return NoteParseResult(
            text="",
            status="empty",
            engine="image-ocr-disabled",
            detail="图片已上传；未配置 OCR（需 QWEN_API_KEY / QWEN_BASE_URL）",
        )

    mime_map = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "webp": "image/webp",
        "gif": "image/gif",
        "avif": "image/avif",
    }
    mime = mime_map.get((ext or "").lower(), "image/png")
    b64 = base64.b64encode(data).decode("ascii")
    image_data_url = f"data:{mime};base64,{b64}"
    prompt = (
        "请对这张图片做 OCR，只输出可读正文，不要解释。"
        "要求：保留段落与换行；忽略装饰元素；若无可读正文返回空字符串。"
    )
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": "你是 OCR 助手，只返回识别后的正文。"},
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": image_data_url}},
            ],
        },
    ]
    try:
        txt = chat_completion_openai_compatible(
            messages=messages,  # type: ignore[arg-type]
            api_base=base,
            api_key=key,
            model=model,
            temperature=0.0,
            timeout_sec=120,
        )
    except Exception as exc:
        logger.warning("image ocr failed model=%s: %s", model, exc)
        return NoteParseResult(
            text="",
            status="empty",
            engine="image-ocr-error",
            detail=f"OCR 失败：{exc}",
        )
    text = str(txt or "").strip()
    if text:
        return NoteParseResult(text=text, status="ok", engine=f"qwen-vl:{model}", detail=None)
    return NoteParseResult(
        text="",
        status="empty",
        engine=f"qwen-vl:{model}",
        detail="OCR 未识别到可用正文",
    )


def extract_pdf_from_bytes(data: bytes) -> NoteParseResult:
    pages, pym_chars = _pdf_pymupdf_chars_per_page(data)
    text, ok = _pdf_pymupdf(data)
    engine = "pymupdf"
    if not ok or len((text or "").strip()) < 8:
        t2, ok2 = _pdf_pypdf2(data)
        if ok2 and len(t2.strip()) > len((text or "").strip()):
            text = t2
            engine = "pypdf2"
            ok = True
        elif not (text or "").strip() and t2.strip():
            text = t2
            engine = "pypdf2"
            ok = True

    stripped = (text or "").strip()
    scanned_detail = (
        "扫描版 PDF 暂不支持：未检测到可用文字层。"
        "请使用 OCR 或本地工具导出为可搜索 PDF / Word 后再上传。"
    )

    if pages >= 1:
        low_layer = pym_chars < max(45, pages * 22)
        if len(stripped) < 40 and low_layer:
            return NoteParseResult(text="", status="empty", engine=engine, detail=scanned_detail)
        if not stripped and low_layer:
            return NoteParseResult(text="", status="empty", engine=engine, detail=scanned_detail)

    if stripped:
        segs = _pdf_rag_segments_by_page(data)
        return NoteParseResult(
            text=stripped,
            status="ok",
            engine=engine,
            detail=None,
            rag_segments=segs if segs else None,
        )
    return NoteParseResult(
        text="",
        status="empty",
        engine=engine,
        detail="无法提取文本（可能为加密 PDF 或扫描件）；加密文件请先解除密码。",
    )


def extract_text_from_bytes(data: bytes, ext: str) -> NoteParseResult:
    """
    从文件字节抽取纯文本（与 parse_note_temp_path / 播客加载共用）。
    """
    e = (ext or "txt").lower().lstrip(".")
    if not data:
        return NoteParseResult(text="", status="empty", engine="none", detail="空文件")

    if e in ("txt", "md", "markdown"):
        text, enc = _decode_plain_bytes(data)
        st = "ok" if text.strip() else "empty"
        return NoteParseResult(
            text=text,
            status=st,
            engine="charset_normalizer" if enc else "utf-8",
            detail=None if st == "ok" else "无可见文本",
            encoding=enc,
        )

    if e == "docx":
        text = _docx_python_docx(data)
        eng = "python-docx"
        segs = _docx_rag_segments(data)
        if not text.strip():
            text = _docx_xml_fallback(data)
            eng = "docx-xml-fallback"
            segs = []
        if text.strip():
            return NoteParseResult(
                text=text,
                status="ok",
                engine=eng,
                detail=None,
                rag_segments=segs if segs else None,
            )
        return NoteParseResult(text="", status="empty", engine=eng, detail="DOCX 未解析出正文")

    if e == "pdf":
        return extract_pdf_from_bytes(data)

    if e == "epub":
        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        try:
            text, ok, segs = _epub_via_content_parser(path)
            if ok and text.strip():
                return NoteParseResult(
                    text=text,
                    status="ok",
                    engine="epub",
                    detail=None,
                    rag_segments=segs,
                )
            return NoteParseResult(
                text="",
                status="empty",
                engine="epub",
                detail="EPUB 未提取到正文",
            )
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass

    if e == "doc":
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        try:
            text, ok = _parse_doc_path(path)
            if ok and text.strip():
                normalized = _normalize_prose_plaintext(text)
                segs = _plain_text_rag_segments(normalized)
                return NoteParseResult(
                    text=normalized,
                    status="ok",
                    engine="antiword|catdoc|soffice",
                    detail=None,
                    rag_segments=segs if segs else None,
                )
            return NoteParseResult(
                text="",
                status="empty",
                engine="doc-binary",
                detail="未安装 antiword/catdoc/LibreOffice 或无法解析该 DOC",
            )
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass

    if e == "csv":
        return _csv_bytes_to_text(data)

    if e == "xlsx":
        return _xlsx_bytes_to_text(data)

    if e == "xls":
        return _xls_bytes_to_text(data)

    if e in ("html", "htm", "xhtml"):
        return _extract_html_from_bytes(data)

    if e in ("png", "jpg", "jpeg", "webp", "gif", "avif"):
        return _ocr_image_via_openai_compat(data, e)

    # 未知扩展名：按纯文本尝试
    text, enc = _decode_plain_bytes(data)
    st = "ok" if text.strip() else "empty"
    return NoteParseResult(
        text=text,
        status=st,
        engine="text-fallback",
        detail=None,
        encoding=enc,
    )


def extract_pdf_dict_for_legacy(pdf_path: str) -> dict[str, Any]:
    """
    与 content_parser.parse_pdf 返回结构兼容，供旧调用方使用。
    """
    try:
        with open(pdf_path, "rb") as f:
            data = f.read()
    except Exception as e:
        return {
            "success": False,
            "error": str(e)[:500],
            "content": "",
            "logs": [],
            "source": "pdf",
        }
    res = extract_pdf_from_bytes(data)
    logs: list[str] = [f"engine={res.engine}", f"status={res.status}"]
    if res.ok:
        return {
            "success": True,
            "content": res.text,
            "logs": logs,
            "source": "pdf",
        }
    return {
        "success": False,
        "error": res.detail or "pdf_empty",
        "content": "",
        "logs": logs,
        "source": "pdf",
    }

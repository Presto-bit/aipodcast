"""HTML / DOCX 上传解析：按段 rag_segments 与段间双换行。"""

from __future__ import annotations

import io

from app.fyv_shared.content_parser import (
    _extract_epub_blocks_from_soup,
    _html_blocks_to_markdown_and_segments,
)
from app.note_document_extract import (
    _docx_rag_segments,
    _extract_html_from_bytes,
    _normalize_prose_plaintext,
    _plain_text_rag_segments,
)
from bs4 import BeautifulSoup


def test_html_blocks_to_markdown_splits_paragraphs() -> None:
    soup = BeautifulSoup(
        "<html><body><h1>标题</h1><p>段落一。</p><p>段落二。</p></body></html>",
        "html.parser",
    )
    blocks = _extract_epub_blocks_from_soup(soup)
    content, segs = _html_blocks_to_markdown_and_segments(blocks, chapter_mode=False)
    assert "# 标题" in content
    assert "段落一" in content and "段落二" in content
    assert "\n\n段落一" in content or "段落一。\n\n段落二" in content
    para = [s for s in segs if (s.get("meta") or {}).get("block_type") == "paragraph"]
    assert len(para) >= 2
    headings = [s for s in segs if (s.get("meta") or {}).get("block_type") == "heading"]
    assert len(headings) >= 1


def test_extract_html_from_bytes_uses_block_parser() -> None:
    html = (
        "<!DOCTYPE html><html><head><title>T</title></head>"
        "<body><h2>小节</h2><p>第一段。</p><p>第二段。</p></body></html>"
    ).encode("utf-8")
    r = _extract_html_from_bytes(html)
    assert r.status == "ok"
    assert r.engine == "html-bs4-blocks"
    assert "## 小节" in (r.text or "")
    assert "\n\n第一段" in (r.text or "") or "第一段。\n\n第二段" in (r.text or "")
    segs = r.rag_segments or []
    assert len([s for s in segs if (s.get("meta") or {}).get("block_type") == "paragraph"]) >= 2


def test_docx_rag_segments_one_paragraph_per_block() -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return
    doc = Document()
    doc.add_paragraph("段落甲。")
    doc.add_paragraph("段落乙。")
    buf = io.BytesIO()
    doc.save(buf)
    segs = _docx_rag_segments(buf.getvalue())
    paras = [s for s in segs if (s.get("meta") or {}).get("block_type") == "paragraph"]
    assert len(paras) == 2
    assert paras[0].get("text") == "段落甲。"
    assert paras[1].get("text") == "段落乙。"


def test_plain_text_rag_segments_from_doc_style_text() -> None:
    raw = "第一行。\n第二行。\n\n第三段。"
    norm = _normalize_prose_plaintext(raw)
    segs = _plain_text_rag_segments(norm)
    assert len(segs) >= 2
